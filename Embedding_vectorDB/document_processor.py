import PyPDF2
from docx import Document
import pandas as pd
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import re
from typing import List, Dict
import cv2
import numpy as np

class DocumentProcessor:
    def __init__(self):
        self.ocr_prompt = """Hãy OCR image này và trả ra text dạng markdown giúp tôi. Nếu có dạng bảng hãy đảm bảo khi OCR hãy tạo ra các cell thật đồng đều cho người khác dễ nhìn. Hãy đảm bảo bạn không gen text trong ảnh. Tôi muốn dữ liệu của tôi, bạn không cần thêm gì cả. Tôi không cần trả lại dữ liệu ảnh, markdown, trích dẫn. Hãy loại bỏ số trang và header/footer không cần thiết"""

    def process_pdf(self, file_path: str) -> str:
        """Process PDF file and convert to continuous markdown without page breaks"""
        markdown_content = ""

        try:
            # First try to extract text directly
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""

                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        # Remove page numbers and common headers/footers
                        page_text = self.clean_page_artifacts(page_text)
                        full_text += f" {page_text}"

                markdown_content = full_text.strip()

            # If no text extracted or minimal text, use OCR
            if len(markdown_content.strip()) < 100:
                markdown_content = self.ocr_pdf_continuous(file_path)

        except Exception as e:
            # Fallback to OCR
            markdown_content = self.ocr_pdf_continuous(file_path)

        return self.clean_and_structure_markdown(markdown_content)

    def ocr_pdf_continuous(self, file_path: str) -> str:
        """OCR PDF pages to extract text as continuous content"""
        full_text = ""

        try:
            # Convert PDF pages to images
            pages = convert_from_path(file_path)

            for page in pages:
                # OCR the page
                page_text = pytesseract.image_to_string(page, lang='vie')
                # Remove page artifacts and add to continuous text
                cleaned_text = self.clean_page_artifacts(page_text)
                if cleaned_text.strip():
                    full_text += f" {cleaned_text}"

        except Exception as e:
            print(f"OCR Error: {e}")

        return full_text.strip()

    def clean_page_artifacts(self, text: str) -> str:
        """Remove page numbers, headers, footers and other page artifacts"""
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            # Skip empty lines
            if not line:
                continue

            # Skip lines that are likely page numbers (just numbers)
            if re.match(r'^\d+$', line):
                continue

            # Skip common header/footer patterns
            if re.match(r'^Trang \d+', line, re.IGNORECASE):
                continue
            if re.match(r'^Page \d+', line, re.IGNORECASE):
                continue
            if len(line) < 5 and line.isdigit():
                continue

            cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    def process_word(self, file_path: str) -> str:
        """Process Word document and convert to continuous markdown"""
        try:
            doc = Document(file_path)
            markdown_content = ""

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Check if it's a heading based on style
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[
                            -1].isdigit() else 1
                        markdown_content += f"\n{'#' * level} {text}\n\n"
                    else:
                        markdown_content += f"{text} "

            # Process tables
            for table in doc.tables:
                markdown_content += "\n\n" + self.convert_table_to_markdown(table) + "\n\n"

            return self.clean_and_structure_markdown(markdown_content)

        except Exception as e:
            raise Exception(f"Error processing Word document: {e}")

    def process_excel(self, file_path: str) -> str:
        """Process Excel file and convert to continuous markdown"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            markdown_content = ""

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                if not df.empty:
                    markdown_content += f"\n## {sheet_name}\n\n"
                    # Convert multi-dimensional table to 1D format
                    markdown_content += self.convert_dataframe_to_1d_markdown(df) + "\n"

            return self.clean_and_structure_markdown(markdown_content)

        except Exception as e:
            raise Exception(f"Error processing Excel file: {e}")

    def process_text(self, text_content: str) -> str:
        """Process plain text and convert to structured markdown"""
        # Remove page breaks and page numbers
        cleaned_text = self.clean_page_artifacts(text_content)
        return self.clean_and_structure_markdown(cleaned_text)

    def convert_dataframe_to_1d_markdown(self, df: pd.DataFrame) -> str:
        """Convert DataFrame to 1D markdown format as requested"""
        markdown = ""

        if df.empty:
            return markdown

        # Get column names
        columns = df.columns.tolist()

        # Convert each row to 1D format
        for index, row in df.iterrows():
            row_text = []
            for col in columns:
                value = str(row[col]) if pd.notna(row[col]) else ""
                if value.strip():
                    row_text.append(f"**{col}**: {value}")

            if row_text:
                markdown += ", ".join(row_text) + "\n\n"

        return markdown

    def convert_table_to_markdown(self, table) -> str:
        """Convert Word table to 1D markdown format"""
        markdown = ""

        # Extract headers from first row
        headers = []
        if table.rows:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip())

        # Process each data row
        for row_idx, row in enumerate(table.rows[1:], 1):
            row_data = []
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(headers):
                    value = cell.text.strip()
                    if value:
                        row_data.append(f"**{headers[col_idx]}**: {value}")

            if row_data:
                markdown += ", ".join(row_data) + "\n\n"

        return markdown

    def clean_and_structure_markdown(self, content: str) -> str:
        """Clean and structure markdown content as continuous text"""
        # Remove excessive whitespace but maintain paragraph structure
        content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)

        # Remove standalone page numbers
        content = re.sub(r'\n\d+\n', '\n', content)
        content = re.sub(r'^Trang \d+.*\n', '', content, flags=re.MULTILINE | re.IGNORECASE)
        content = re.sub(r'^Page \d+.*\n', '', content, flags=re.MULTILINE | re.IGNORECASE)

        # Process lines for better structure
        lines = content.split('\n')
        structured_lines = []
        previous_was_heading = False

        for line in lines:
            line = line.strip()
            if not line:
                if not previous_was_heading:
                    structured_lines.append('')
                continue

            # Auto-detect headings (lines that are all caps or start with numbers)
            is_heading = False
            if (line.isupper() and len(line.split()) <= 10 and len(line.split()) > 1) or \
                    re.match(r'^\d+\.?\s+[A-ZÀ-Ỹ]', line):
                if not line.startswith('#'):
                    line = f"## {line}"
                    is_heading = True

            structured_lines.append(line)
            previous_was_heading = is_heading

        # Join and clean up final content
        final_content = '\n'.join(structured_lines)

        # Remove multiple consecutive newlines
        final_content = re.sub(r'\n{3,}', '\n\n', final_content)

        return final_content.strip()

    def parse_markdown_to_sentences(self, markdown_content: str) -> List[Dict]:
        """Parse markdown into 3-sentence chunks with section context"""
        sentence_chunks = []
        lines = markdown_content.split('\n')

        current_section = ""
        current_content = ""
        section_hierarchy = []

        def split_into_sentences(text: str) -> List[str]:
            """Split text into sentences using regex patterns"""
            # Cải tiến pattern để xử lý tốt hơn tiếng Việt và tiếng Anh
            sentence_endings = r'[.!?]+'

            # Tách câu dựa trên dấu chấm, chấm hỏi, chấm than
            sentences = re.split(f'({sentence_endings})', text)

            result_sentences = []
            current_sentence = ""

            for i in range(0, len(sentences), 2):
                if i < len(sentences):
                    sentence_text = sentences[i].strip()
                    ending = sentences[i + 1] if i + 1 < len(sentences) else ""

                    if sentence_text:
                        full_sentence = (sentence_text + ending).strip()
                        if full_sentence and len(full_sentence) > 10:  # Lọc câu quá ngắn
                            result_sentences.append(full_sentence)

            return result_sentences

        def process_content_to_sentences(section_title: str, content: str):
            """Process content and create 3-sentence chunks"""
            if not content.strip():
                return

            # Tách nội dung thành câu
            sentences = split_into_sentences(content)

            # Nhóm các câu thành chunks 3 câu
            for i in range(0, len(sentences), 3):
                # Lấy tối đa 3 câu cho mỗi chunk
                chunk_sentences = sentences[i:i + 3]

                if chunk_sentences:  # Đảm bảo chunk không rỗng
                    # Kết hợp các câu trong chunk
                    combined_sentences = ' '.join(chunk_sentences)

                    sentence_chunks.append({
                        'section_title': section_title,
                        'content': f"{section_title}\n{combined_sentences}",
                        'sentences': chunk_sentences,  # Danh sách các câu riêng lẻ
                        'combined_text': combined_sentences,  # Text đã kết hợp
                        'sentence_count': len(chunk_sentences),  # Số câu trong chunk
                        'context': section_title
                    })

        for line in lines:
            line = line.strip()

            # Check if line is a heading
            if line.startswith('#'):
                # Process previous content into sentences before moving to new section
                if current_content.strip():
                    process_content_to_sentences(current_section, current_content)
                    current_content = ""

                # Update section hierarchy
                level = len(line.split()[0])  # Count # characters
                heading_text = line[level:].strip()

                # Update hierarchy
                if level <= len(section_hierarchy):
                    section_hierarchy = section_hierarchy[:level - 1]
                section_hierarchy.append(heading_text)

                current_section = '\n'.join([f"{'#' * (i + 1)} {title}" for i, title in enumerate(section_hierarchy)])

            else:
                if line:  # Non-empty content line
                    current_content += f"{line}\n"

        # Process final content into sentences
        if current_content.strip():
            process_content_to_sentences(current_section, current_content)

        return sentence_chunks

        # Phiên bản tối ưu hóa cho tiếng Việt

    def parse_markdown_to_sentences_vietnamese(self, markdown_content: str) -> List[Dict]:
        """Parse markdown into sentence-level chunks optimized for Vietnamese text"""
        sentence_chunks = []
        lines = markdown_content.split('\n')

        current_section = ""
        current_content = ""
        section_hierarchy = []

        def split_vietnamese_sentences(text: str) -> List[str]:
            """Split Vietnamese text into sentences with better accuracy"""
            # Patterns cho tiếng Việt
            patterns = [
                r'[.!?]+\s+',  # Dấu câu + khoảng trắng
                r'[.!?]+\n',  # Dấu câu + xuống dòng
                r'[.!?]+$',  # Dấu câu cuối text
            ]

            # Tách câu
            sentences = []
            current_pos = 0

            for match in re.finditer('|'.join(patterns), text):
                sentence = text[current_pos:match.end()].strip()
                if sentence and len(sentence) > 15:  # Lọc câu quá ngắn
                    sentences.append(sentence)
                current_pos = match.end()

            # Thêm phần còn lại nếu có
            remaining = text[current_pos:].strip()
            if remaining and len(remaining) > 15:
                sentences.append(remaining)

            return sentences

        def process_content_to_sentences(section_title: str, content: str):
            """Process content and create sentence-level chunks"""
            if not content.strip():
                return

            sentences = split_vietnamese_sentences(content)

            for sentence in sentences:
                sentence_clean = sentence.strip()
                if sentence_clean:
                    sentence_chunks.append({
                        'section_title': section_title,
                        'content': f"{section_title}\n{sentence_clean}",
                        'sentence': sentence_clean,
                        'context': section_title,
                        'length': len(sentence_clean)
                    })

        for line in lines:
            line = line.strip()

            if line.startswith('#'):
                # Process previous content
                if current_content.strip():
                    process_content_to_sentences(current_section, current_content)
                    current_content = ""

                # Update section hierarchy
                level = len(line.split()[0])
                heading_text = line[level:].strip()

                if level <= len(section_hierarchy):
                    section_hierarchy = section_hierarchy[:level - 1]
                section_hierarchy.append(heading_text)

                current_section = '\n'.join([f"{'#' * (i + 1)} {title}" for i, title in enumerate(section_hierarchy)])

            else:
                if line:
                    current_content += f"{line}\n"

        # Process final content
        if current_content.strip():
            process_content_to_sentences(current_section, current_content)

        return sentence_chunks